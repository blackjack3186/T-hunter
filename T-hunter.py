#########################################################################
#--------------------TAUTOLOGY---HUNTER---1.1---------------------------#
#---------------------------BY--ESQUELLO--------------------------------#
#########################################################################

import subprocess
import sys
import os
import re
import threading
import time
from datetime import datetime

# --- Установка зависимостей ---
def install_and_import(package, import_name):
    try:
        __import__(import_name)
    except ImportError:
        print(f"Устанавливаю {package}...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])
        os.execl(sys.executable, sys.executable, *sys.argv)

for pkg, imp in [
    ("pyperclip", "pyperclip"),
    ("pymorphy3", "pymorphy3"),
    ("python-docx", "docx"),
]:
    install_and_import(pkg, imp)

import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox
import pyperclip
import pymorphy3
from docx import Document as DocxDocument
from docx.shared import RGBColor


morph = pymorphy3.MorphAnalyzer()

# ===================== НАСТРОЙКИ =====================

WINDOW = 3
MIN_WORD_LEN = 4
MARKER = "*"

PAIR_COLORS = [
    "#CC0000", "#0066CC", "#CC6600", "#008844", "#AA00AA",
    "#CC0066", "#4488AA", "#886600", "#6644CC", "#AA4400",
]
WEED_COLOR = "#8800AA"

STOP_WORDS = {
    "быть", "мочь", "этот", "свой", "весь", "который",
    "один", "такой", "сказать", "говорить", "очень",
    "ещё", "еще", "только", "также", "однако", "когда",
    "потом", "после", "перед", "через", "между", "было",
    "будет", "если", "чтобы", "более", "можно", "нужно",
    "здесь", "туда", "куда", "откуда", "тоже", "себя",
}

WEEDS = [
    "юноша", "побелевшие костяшки", "похолодел",
    "звенящая тишина", "кивнул", "усмехнулся",
    "нахмурился", "вздохнул", "пробормотал",
]

# Копия дефолтных сорняков для кнопки «сброс»
DEFAULT_WEEDS = list(WEEDS)

# Файл для хранения пользовательского списка
WEEDS_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "weeds.txt")

def load_weeds_from_file():
    global WEEDS
    if os.path.exists(WEEDS_FILE):
        with open(WEEDS_FILE, "r", encoding="utf-8") as f:
            loaded = [line.strip() for line in f if line.strip()]
            if loaded:
                WEEDS.clear()
                WEEDS.extend(loaded)

def save_weeds_to_file():
    with open(WEEDS_FILE, "w", encoding="utf-8") as f:
        f.write("\n".join(WEEDS))

# Загружаем при старте
load_weeds_from_file()

# ===================== АНАЛИЗ =====================

def get_lemma(word):
    return morph.parse(word)[0].normal_form

def get_stem(word):
    lemma = get_lemma(word)
    cut = max(3, int(len(lemma) * 0.7))
    return lemma[:cut]

def split_sentences(text):
    sents = re.split(r'(?<=[.!?…»])\s+', text)
    return [s.strip() for s in sents if s.strip()]

def tokenize(sentence):
    return re.findall(r'[а-яёА-ЯЁ]+', sentence.lower())

def analyze(text):
    sentences = split_sentences(text)
    issues = []
    mark_map = {}
    pair_counter = 0

    for i, sent in enumerate(sentences):
        words_i = tokenize(sent)
        lemmas_i = {}
        stems_i = {}
        for w in words_i:
            if len(w) < MIN_WORD_LEN:
                continue
            lem = get_lemma(w)
            if lem in STOP_WORDS:
                continue
            lemmas_i[w] = lem
            stems_i[w] = get_stem(w)

        start = max(0, i - WINDOW)
        end = min(len(sentences), i + WINDOW + 1)

        for j in range(start, end):
            if j <= i:
                continue
            words_j = tokenize(sentences[j])
            for nw in words_j:
                if len(nw) < MIN_WORD_LEN:
                    continue
                n_lem = get_lemma(nw)
                if n_lem in STOP_WORDS:
                    continue
                n_stem = get_stem(nw)

                for cw, c_lem in lemmas_i.items():
                    match_type = None
                    if c_lem == n_lem:
                        match_type = "ПОВТОР"
                        info = c_lem
                    elif stems_i[cw] == n_stem:
                        match_type = "ОДНОКОР"
                        info = f"{c_lem}/{n_lem}"

                    if match_type:
                        key1 = (i, cw)
                        key2 = (j, nw)
                        existing = mark_map.get(key1) or mark_map.get(key2)
                        if existing is not None:
                            pid = existing
                        else:
                            pid = pair_counter
                            pair_counter += 1

                        mark_map[key1] = pid
                        mark_map[key2] = pid

                        issues.append({
                            "type": match_type,
                            "w1": cw, "w2": nw,
                            "info": info,
                            "s1": i + 1, "s2": j + 1,
                            "ctx": sent[:100],
                            "pair_id": pid,
                        })

    for i, sent in enumerate(sentences):
        lower = sent.lower()
        for weed in WEEDS:
            if weed.lower() in lower:
                issues.append({
                    "type": "СОРНЯК",
                    "w1": weed, "w2": "",
                    "info": weed,
                    "s1": i + 1, "s2": 0,
                    "ctx": sent[:100],
                    "pair_id": -1,
                })

    seen = set()
    unique = []
    for iss in issues:
        if iss["type"] == "СОРНЯК":
            key = (iss["type"], iss["info"], iss["s1"])
        else:
            pair = tuple(sorted([iss["s1"], iss["s2"]]))
            key = (iss["type"], iss["info"], pair)
        if key not in seen:
            seen.add(key)
            unique.append(iss)

    return sentences, sorted(unique, key=lambda x: x["s1"]), mark_map


# ===================== GUI =====================

class TautologyHunterApp:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Tautology Hunter 1.1")
        self.root.geometry("620x700")
        self.root.resizable(True, True)
        self.root.minsize(500, 500)
        self.root.protocol("WM_DELETE_WINDOW", self.quit_app)

        self.auto_watch = tk.BooleanVar(value=False)
        self.prev_text = ""
        self.stop_event = threading.Event()
        self.last_marked = ""
        self.last_sentences = []
        self.last_issues = []
        self.last_mark_map = {}

        self._build_ui()

        self.watcher_thread = threading.Thread(
            target=self._clipboard_watcher, daemon=True
        )
        self.watcher_thread.start()

    def _build_ui(self):
        settings = ttk.LabelFrame(self.root, text="  Настройки  ")
        settings.pack(fill="x", padx=10, pady=(10, 5))

        ttk.Checkbutton(
            settings,
            text="Автоматически проверять при копировании",
            variable=self.auto_watch,
        ).pack(anchor="w", padx=10, pady=2)

        btn_frame = ttk.Frame(self.root)
        btn_frame.pack(fill="x", padx=10, pady=5)

        ttk.Button(
            btn_frame, text="🔍 Проверить буфер",
            command=self.check_clipboard,
        ).pack(fill="x", pady=2)

        ttk.Button(
            btn_frame, text="📋 Вставить помеченный текст в буфер",
            command=self.copy_marked,
        ).pack(fill="x", pady=2)

        ttk.Button(
            btn_frame, text="⚙ Настроить слова-сорняки",
            command=self.edit_weeds,
        ).pack(fill="x", pady=2)
        
        ttk.Button(
            btn_frame, text="📄 Проверить DOCX-файл",
            command=self.analyze_docx_file,
        ).pack(fill="x", pady=2)
        self.mark_mode = tk.StringVar(value="stars")

        mode_frame = ttk.LabelFrame(self.root, text="  Режим пометки (beta)  ")
        mode_frame.pack(fill="x", padx=10, pady=5)

        ttk.Radiobutton(
            mode_frame, text="Пометить звёздочкой *",
            variable=self.mark_mode, value="stars",
        ).pack(anchor="w", padx=10, pady=2)

        ttk.Radiobutton(
            mode_frame, text="Создать DOCX-файл",
            variable=self.mark_mode, value="docx",
        ).pack(anchor="w", padx=10, pady=2)
		
        freq_frame = ttk.LabelFrame(self.root, text="  Частотный анализ  ")
        freq_frame.pack(fill="x", padx=10, pady=5)

        freq_settings = ttk.Frame(freq_frame)
        freq_settings.pack(fill="x", padx=10, pady=5)

        ttk.Label(freq_settings, text="Мин. букв:").pack(side="left")
        self.freq_min_len = tk.IntVar(value=4)
        ttk.Spinbox(
            freq_settings, from_=2, to=15, width=4,
            textvariable=self.freq_min_len,
        ).pack(side="left", padx=(2, 15))

        ttk.Label(freq_settings, text="Мин. повторов:").pack(side="left")
        self.freq_min_count = tk.IntVar(value=3)
        ttk.Spinbox(
            freq_settings, from_=2, to=100, width=4,
            textvariable=self.freq_min_count,
        ).pack(side="left", padx=(2, 15))

        ttk.Button(
            freq_frame, text="📊 Анализ частот из файла",
            command=self.run_frequency_analysis,
        ).pack(fill="x", padx=10, pady=(0, 5))
        
        ttk.Separator(freq_frame, orient="horizontal").pack(fill="x", padx=10, pady=5)

        ngram_settings = ttk.Frame(freq_frame)
        ngram_settings.pack(fill="x", padx=10, pady=5)

        ttk.Label(ngram_settings, text="Слов в фразе:").pack(side="left")
        self.ngram_size = tk.IntVar(value=2)
        ttk.Spinbox(
            ngram_settings, from_=2, to=5, width=4,
            textvariable=self.ngram_size,
        ).pack(side="left", padx=(2, 15))

        ttk.Label(ngram_settings, text="Мин. повторов:").pack(side="left")
        self.ngram_min_count = tk.IntVar(value=2)
        ttk.Spinbox(
            ngram_settings, from_=2, to=50, width=4,
            textvariable=self.ngram_min_count,
        ).pack(side="left", padx=(2, 15))

        ttk.Button(
            freq_frame, text="🔗 Повторяющиеся фразы из файла",
            command=self.run_ngram_analysis,
        ).pack(fill="x", padx=10, pady=(0, 5))

        results_frame = ttk.LabelFrame(self.root, text="  Найденные проблемы  ")
        results_frame.pack(fill="both", expand=True, padx=10, pady=5)

        self.results_text = scrolledtext.ScrolledText(
            results_frame, height=12, wrap="word",
            font=("Consolas", 9), state="disabled",
        )
        self.results_text.pack(fill="both", expand=True, padx=5, pady=5)
        self.results_text.tag_configure("weed", foreground=WEED_COLOR)

        preview_frame = ttk.LabelFrame(self.root, text="  Помеченный текст (превью)  ")
        preview_frame.pack(fill="both", expand=True, padx=10, pady=(5, 10))

        self.preview_text = scrolledtext.ScrolledText(
            preview_frame, height=8, wrap="word",
            font=("Consolas", 9), state="disabled",
        )
        self.preview_text.pack(fill="both", expand=True, padx=5, pady=5)

        self.status_var = tk.StringVar(value="Готов. Скопируй текст и нажми «Проверить».")
        ttk.Label(
            self.root, textvariable=self.status_var, foreground="gray"
        ).pack(pady=(0, 5))
        
    def check_clipboard(self):
        text = pyperclip.paste()
        if not text or len(text.strip()) < 20:
            messagebox.showwarning("Пусто", "Скопируй текст подлиннее.")
            return

        self.status_var.set("Анализирую...")
        self.root.update()

        try:
            sentences, issues, mark_map = analyze(text)
        except Exception as e:
            messagebox.showerror("Ошибка", str(e))
            self.status_var.set("Ошибка анализа.")
            return

        self.last_sentences = sentences
        self.last_issues = issues
        self.last_mark_map = mark_map
        self._build_marked_text()
        self._show_results(issues)
        self._show_preview()

        repeats = sum(1 for i in issues if i["type"] == "ПОВТОР")
        cognates = sum(1 for i in issues if i["type"] == "ОДНОКОР")
        weeds = sum(1 for i in issues if i["type"] == "СОРНЯК")
        self.status_var.set(
            f"Повторов: {repeats} | Однокоренных: {cognates} | Сорняков: {weeds}"
        )

    def _build_marked_text(self):
        original = pyperclip.paste()
        paragraphs = original.split("\n")

        # Собираем леммы для пометки
        lemma_to_mark = set()
        for (si, w), pid in self.last_mark_map.items():
            lemma_to_mark.add(get_lemma(w))

        weed_set = {w.lower() for w in WEEDS}

        marked_paras = []
        for para_text in paragraphs:
            if not para_text.strip():
                marked_paras.append("")
                continue

            def mark_word(m):
                low = m.group(0).lower()
                if low in weed_set:
                    return f"{MARKER}{m.group(0)}{MARKER}"
                if len(low) >= MIN_WORD_LEN:
                    lem = get_lemma(low)
                    if lem in lemma_to_mark:
                        return f"{MARKER}{m.group(0)}{MARKER}"
                return m.group(0)

            result = re.sub(r'[а-яёА-ЯЁ]+', mark_word, para_text)

            # Многословные сорняки
            for weed in WEEDS:
                if " " in weed and weed.lower() in result.lower():
                    pattern = re.compile(re.escape(weed), re.IGNORECASE)
                    result = pattern.sub(f"{MARKER}{weed}{MARKER}", result)

            marked_paras.append(result)

        self.last_marked = "\n".join(marked_paras)

    def _show_results(self, issues):
        self.results_text.configure(state="normal")
        self.results_text.delete("1.0", "end")

        if not issues:
            self.results_text.insert("end", "Ничего не найдено. Чистый текст!\n")
        else:
            for iss in issues:
                if iss["type"] == "СОРНЯК":
                    tag = "weed"
                else:
                    color_idx = iss["pair_id"] % len(PAIR_COLORS)
                    tag = f"pair_{iss['pair_id']}"
                    self.results_text.tag_configure(
                        tag,
                        foreground=PAIR_COLORS[color_idx],
                        font=("Consolas", 9, "bold"),
                    )

                if iss["type"] == "СОРНЯК":
                    line = f"[{iss['type']}] «{iss['w1']}» — предл. {iss['s1']}\n"
                else:
                    line = (
                        f"[{iss['type']}] «{iss['w1']}» ↔ «{iss['w2']}» "
                        f"({iss['info']}) — предл. {iss['s1']} и {iss['s2']}\n"
                    )

                self.results_text.insert("end", line, tag)
                self.results_text.insert("end", f"   → {iss['ctx']}…\n\n")

        self.results_text.configure(state="disabled")

    def _show_preview(self):
        self.preview_text.configure(state="normal")
        self.preview_text.delete("1.0", "end")

        for (si, w), pid in self.last_mark_map.items():
            tag = f"preview_pair_{pid}"
            color = PAIR_COLORS[pid % len(PAIR_COLORS)]
            self.preview_text.tag_configure(
                tag, foreground=color, font=("Consolas", 9, "bold")
            )
        self.preview_text.tag_configure(
            "preview_weed",
            foreground=WEED_COLOR,
            font=("Consolas", 9, "bold", "underline"),
        )

        for i, sent in enumerate(self.last_sentences):
            words_in_sent = {
                w: pid for (si, w), pid in self.last_mark_map.items() if si == i
            }
            weed_words = set()
            for weed in WEEDS:
                if weed.lower() in sent.lower():
                    weed_words.add(weed.lower())

            tokens = re.split(r'([а-яёА-ЯЁ]+)', sent)
            for token in tokens:
                low = token.lower()
                if low in weed_words:
                    self.preview_text.insert("end", token, "preview_weed")
                elif low in words_in_sent:
                    pid = words_in_sent[low]
                    self.preview_text.insert("end", token, f"preview_pair_{pid}")
                else:
                    self.preview_text.insert("end", token)

            self.preview_text.insert("end", " ")

        self.preview_text.configure(state="disabled")

    def copy_marked(self):
        if not self.last_sentences:
            messagebox.showinfo("Пусто", "Сначала проверь текст.")
            return

        if self.mark_mode.get() == "stars":
            pyperclip.copy(self.last_marked)
            self.prev_text = self.last_marked
            self.status_var.set("Помеченный текст скопирован в буфер.")
        else:
            self._save_colored_docx()

    # --- Автонаблюдение ---

    def _clipboard_watcher(self):
        while not self.stop_event.is_set():
            try:
                if self.auto_watch.get():
                    text = pyperclip.paste()
                    if text and text != self.prev_text and len(text) > 50:
                        self.prev_text = text
                        self.root.after(0, self.check_clipboard)
            except Exception:
                pass
            time.sleep(1)

    # --- Lifecycle ---

    def quit_app(self):
        self.stop_event.set()
        self.root.destroy()

    def run(self):
        self.root.mainloop()
        
    def edit_weeds(self):
        win = tk.Toplevel(self.root)
        win.title("Слова-сорняки")
        win.geometry("420x500")
        win.resizable(False, False)
        win.attributes("-topmost", True)
        win.grab_set()

        tk.Label(
            win, text="Каждое слово или фраза — на отдельной строке.",
            foreground="gray",
        ).pack(padx=10, pady=(10, 2), anchor="w")

        text_input = scrolledtext.ScrolledText(
            win, height=18, wrap="word", font=("Consolas", 10)
        )
        text_input.pack(fill="both", expand=True, padx=10, pady=5)
        text_input.insert("1.0", "\n".join(WEEDS))

        btn_row = ttk.Frame(win)
        btn_row.pack(fill="x", padx=10, pady=(0, 10))

        def save():
            global WEEDS
            raw = text_input.get("1.0", "end").strip()
            new_weeds = [line.strip() for line in raw.split("\n") if line.strip()]
            WEEDS.clear()
            WEEDS.extend(new_weeds)
            save_weeds_to_file()
            self.status_var.set(f"Сорняков в списке: {len(WEEDS)}")
            win.destroy()

        def reset():
            text_input.delete("1.0", "end")
            text_input.insert("1.0", "\n".join(DEFAULT_WEEDS))

        ttk.Button(btn_row, text="✅ Сохранить", command=save).pack(
            side="right", padx=(5, 0)
        )
        ttk.Button(btn_row, text="🔄 По умолчанию", command=reset).pack(
            side="right"
        )
        
    def _save_colored_docx(self):
        from tkinter import filedialog

        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word", "*.docx")],
            title="Сохранить размеченный файл",
        )
        if not path:
            return

        doc = DocxDocument()

        def hex_to_rgb(h):
            h = h.lstrip("#")
            return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

        pair_rgbs = [hex_to_rgb(c) for c in PAIR_COLORS]
        weed_rgb = hex_to_rgb(WEED_COLOR)

        # Собираем словарь: лемма -> pair_id (из всех найденных пар)
        lemma_to_pid = {}
        for (si, w), pid in self.last_mark_map.items():
            lem = get_lemma(w)
            lemma_to_pid[lem] = pid

        weed_set = {w.lower() for w in WEEDS}

        # Берём оригинальный текст и разбиваем по абзацам
        original = pyperclip.paste()
        paragraphs = original.split("\n")

        for para_text in paragraphs:
            para = doc.add_paragraph()

            if not para_text.strip():
                continue

            tokens = re.split(r'([а-яёА-ЯЁ]+)', para_text)
            for token in tokens:
                run = para.add_run(token)
                low = token.lower()

                # Проверяем сорняки
                is_weed = any(w in low for w in weed_set if len(w.split()) == 1)
                if is_weed and low in weed_set:
                    run.font.color.rgb = weed_rgb
                    run.bold = True
                    run.underline = True
                    continue

                # Проверяем повторы/однокоренные
                if len(low) >= MIN_WORD_LEN and re.match(r'^[а-яёА-ЯЁ]+$', low):
                    lem = get_lemma(low)
                    if lem in lemma_to_pid:
                        pid = lemma_to_pid[lem]
                        run.font.color.rgb = pair_rgbs[pid % len(pair_rgbs)]
                        run.bold = True

        try:
            doc.save(path)
            self.status_var.set(f"Сохранено: {os.path.basename(path)}")
            if sys.platform == "win32":
                os.startfile(path)
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить:\n{e}")

    def run_frequency_analysis(self):
        from tkinter import filedialog
        import csv

        # Выбор входного файла
        input_path = filedialog.askopenfilename(
            filetypes=[
                ("Word", "*.docx"),
                ("Текст", "*.txt"),
            ],
            title="Выбери файл для анализа",
        )
        if not input_path:
            return

        self.status_var.set("Считаю частоты...")
        self.root.update()

        # Извлекаем текст
        try:
            if input_path.endswith(".docx"):
                doc = DocxDocument(input_path)
                full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            else:
                with open(input_path, "r", encoding="utf-8") as f:
                    full_text = f.read()
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось прочитать файл:\n{e}")
            self.status_var.set("Ошибка чтения.")
            return

        # Считаем леммы
        words = re.findall(r'[а-яёА-ЯЁ]+', full_text.lower())
        min_len = self.freq_min_len.get()
        min_count = self.freq_min_count.get()

        from collections import Counter
        lemma_counter = Counter()
        lemma_examples = {}

        for w in words:
            if len(w) < min_len:
                continue
            lem = get_lemma(w)
            if lem in STOP_WORDS:
                continue
            lemma_counter[lem] += 1
            if lem not in lemma_examples:
                lemma_examples[lem] = w

        # Фильтруем
        frequent = [
            (lem, count, lemma_examples[lem])
            for lem, count in lemma_counter.most_common()
            if count >= min_count
        ]

        if not frequent:
            messagebox.showinfo(
                "Результат",
                f"Слов с {min_count}+ повторами (от {min_len} букв) не найдено.",
            )
            self.status_var.set("Частых слов не найдено.")
            return

        # Показываем результат в окне
        win = tk.Toplevel(self.root)
        win.title(f"Частотный анализ — {len(frequent)} слов")
        win.geometry("500x500")
        win.resizable(True, True)
        win.attributes("-topmost", True)

        # Таблица
        columns = ("lemma", "example", "count")
        tree = ttk.Treeview(win, columns=columns, show="headings", height=20)
        tree.heading("lemma", text="Лемма")
        tree.heading("example", text="Пример")
        tree.heading("count", text="Кол-во")
        tree.column("lemma", width=180)
        tree.column("example", width=180)
        tree.column("count", width=80, anchor="center")

        scrollbar = ttk.Scrollbar(win, orient="vertical", command=tree.yview)
        tree.configure(yscrollcommand=scrollbar.set)

        tree.pack(side="left", fill="both", expand=True, padx=(10, 0), pady=10)
        scrollbar.pack(side="left", fill="y", pady=10, padx=(0, 5))

        for lem, count, example in frequent:
            tree.insert("", "end", values=(lem, example, count))

        btn_row = ttk.Frame(win)
        btn_row.pack(fill="x", padx=10, pady=(0, 10))

        def save_csv():
            csv_path = filedialog.asksaveasfilename(
                defaultextension=".csv",
                filetypes=[("CSV", "*.csv")],
                title="Сохранить CSV",
            )
            if not csv_path:
                return
            try:
                with open(csv_path, "w", encoding="utf-8-sig", newline="") as f:
                    writer = csv.writer(f, delimiter=";")
                    writer.writerow(["Лемма", "Пример", "Количество"])
                    for lem, count, example in frequent:
                        writer.writerow([lem, example, count])
                self.status_var.set(f"CSV сохранён: {os.path.basename(csv_path)}")
                if sys.platform == "win32":
                    os.startfile(csv_path)
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось сохранить:\n{e}")

        def add_to_weeds():
            selected = tree.selection()
            if not selected:
                messagebox.showinfo("Выбери", "Выдели строки, которые хочешь добавить в сорняки.")
                return
            added = 0
            for item in selected:
                vals = tree.item(item, "values")
                lemma = vals[0]
                if lemma not in WEEDS:
                    WEEDS.append(lemma)
                    added += 1
            if added:
                save_weeds_to_file()
                self.status_var.set(f"Добавлено сорняков: {added}")

        ttk.Button(btn_row, text="💾 Сохранить CSV", command=save_csv).pack(
            side="left", padx=(0, 5)
        )
        ttk.Button(btn_row, text="🚫 Добавить в сорняки", command=add_to_weeds).pack(
            side="left"
        )
    
    def run_ngram_analysis(self):
        from tkinter import filedialog
        import csv
        from collections import Counter

        input_path = filedialog.askopenfilename(
            filetypes=[
                ("Word", "*.docx"),
                ("Текст", "*.txt"),
            ],
            title="Выбери файл для анализа фраз",
        )
        if not input_path:
            return

        self.status_var.set("Ищу повторяющиеся фразы...")
        self.root.update()

        # Извлекаем текст
        try:
            if input_path.endswith(".docx"):
                doc = DocxDocument(input_path)
                full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            else:
                with open(input_path, "r", encoding="utf-8") as f:
                    full_text = f.read()
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось прочитать файл:\n{e}")
            self.status_var.set("Ошибка чтения.")
            return

        # Разбиваем на предложения, потом на леммы
        sentences = split_sentences(full_text)
        n = self.ngram_size.get()
        min_count = self.ngram_min_count.get()

        ngram_counter = Counter()
        ngram_originals = {}  # лемма-кортеж -> пример из текста

        for sent in sentences:
            words = re.findall(r'[а-яёА-ЯЁ]+', sent)
            if len(words) < n:
                continue

            # Леммы и оригиналы
            lemmas = []
            originals = []
            for w in words:
                lem = get_lemma(w.lower())
                lemmas.append(lem)
                originals.append(w)

            # Собираем n-граммы
            for i in range(len(lemmas) - n + 1):
                ngram = tuple(lemmas[i:i + n])

                # Пропускаем, если все слова короткие или стоп-слова
                meaningful = [w for w in ngram if len(w) >= 3 and w not in STOP_WORDS]
                if len(meaningful) < 2:
                    continue

                ngram_counter[ngram] += 1

                if ngram not in ngram_originals:
                    ngram_originals[ngram] = " ".join(originals[i:i + n])

        # Фильтруем
        frequent = [
            (ngram, count, ngram_originals[ngram])
            for ngram, count in ngram_counter.most_common()
            if count >= min_count
        ]

        if not frequent:
            messagebox.showinfo(
                "Результат",
                f"Фраз из {n} слов с {min_count}+ повторами не найдено.",
            )
            self.status_var.set("Повторяющихся фраз не найдено.")
            return

        # Окно результатов
        win = tk.Toplevel(self.root)
        win.title(f"Повторяющиеся фразы ({n} сл.) — {len(frequent)} шт.")
        win.geometry("600x500")
        win.resizable(True, True)
        win.attributes("-topmost", True)

        columns = ("phrase", "example", "count")
        tree = ttk.Treeview(win, columns=columns, show="headings", height=20)
        tree.heading("phrase", text="Лемма-фраза")
        tree.heading("example", text="Пример из текста")
        tree.heading("count", text="Кол-во")
        tree.column("phrase", width=200)
        tree.column("example", width=250)
        tree.column("count", width=80, anchor="center")

        scrollbar = ttk.Scrollbar(win, orient="vertical", command=tree.yview)
        tree.configure(yscrollcommand=scrollbar.set)

        tree.pack(side="left", fill="both", expand=True, padx=(10, 0), pady=10)
        scrollbar.pack(side="left", fill="y", pady=10, padx=(0, 5))

        for ngram, count, example in frequent:
            phrase_str = " + ".join(ngram)
            tree.insert("", "end", values=(phrase_str, example, count))

        btn_row = ttk.Frame(win)
        btn_row.pack(fill="x", padx=10, pady=(0, 10))

        def save_csv():
            csv_path = filedialog.asksaveasfilename(
                defaultextension=".csv",
                filetypes=[("CSV", "*.csv")],
                title="Сохранить CSV",
            )
            if not csv_path:
                return
            try:
                with open(csv_path, "w", encoding="utf-8-sig", newline="") as f:
                    writer = csv.writer(f, delimiter=";")
                    writer.writerow(["Лемма-фраза", "Пример", "Количество"])
                    for ngram, count, example in frequent:
                        writer.writerow([" + ".join(ngram), example, count])
                self.status_var.set(f"CSV сохранён: {os.path.basename(csv_path)}")
                if sys.platform == "win32":
                    os.startfile(csv_path)
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось сохранить:\n{e}")

        def add_to_weeds():
            selected = tree.selection()
            if not selected:
                messagebox.showinfo("Выбери", "Выдели фразы для добавления в сорняки.")
                return
            added = 0
            for item in selected:
                vals = tree.item(item, "values")
                example = vals[1]
                if example not in WEEDS:
                    WEEDS.append(example.lower())
                    added += 1
            if added:
                save_weeds_to_file()
                self.status_var.set(f"Добавлено сорняков: {added}")

        ttk.Button(btn_row, text="💾 Сохранить CSV", command=save_csv).pack(
            side="left", padx=(0, 5)
        )
        ttk.Button(btn_row, text="🚫 Добавить в сорняки", command=add_to_weeds).pack(
            side="left"
        )

        self.status_var.set(f"Найдено фраз: {len(frequent)}")

    def analyze_docx_file(self):
        from tkinter import filedialog
        from docx.shared import Pt
        from copy import deepcopy

        input_path = filedialog.askopenfilename(
            filetypes=[("Word", "*.docx")],
            title="Выбери DOCX для проверки",
        )
        if not input_path:
            return

        self.status_var.set("Анализирую файл...")
        self.root.update()

        try:
            doc = DocxDocument(input_path)
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось открыть:\n{e}")
            return

        # Извлекаем текст и анализируем
        full_text = "\n".join(p.text for p in doc.paragraphs)
        sentences, issues, mark_map = analyze(full_text)

        if not issues:
            messagebox.showinfo("Чисто!", "Проблем не найдено.")
            self.status_var.set("Файл чист.")
            return

        # Собираем леммы для подсветки
        lemma_to_pid = {}
        for (si, w), pid in mark_map.items():
            lem = get_lemma(w)
            lemma_to_pid[lem] = pid

        weed_set = {w.lower() for w in WEEDS}

        def hex_to_rgb(h):
            h = h.lstrip("#")
            return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

        pair_rgbs = [hex_to_rgb(c) for c in PAIR_COLORS]
        weed_rgb = hex_to_rgb(WEED_COLOR)

        # Создаём новый документ с подсветкой
        new_doc = DocxDocument()

        # Копируем стиль шрифта из оригинала (если есть)
        default_size = Pt(12)

        for para in doc.paragraphs:
            new_para = new_doc.add_paragraph()

            # Копируем выравнивание
            if para.alignment is not None:
                new_para.alignment = para.alignment

            # Копируем стиль абзаца
            try:
                new_para.style = para.style
            except Exception:
                pass

            if not para.text.strip():
                continue

            tokens = re.split(r'([а-яёА-ЯЁ]+)', para.text)

            for token in tokens:
                run = new_para.add_run(token)
                low = token.lower()

                # Восстанавливаем размер шрифта
                run.font.size = default_size

                # Проверяем сорняки (однословные)
                if low in weed_set:
                    run.font.color.rgb = weed_rgb
                    run.bold = True
                    run.underline = True
                    continue

                # Проверяем многословные сорняки — пометим позже
                # Проверяем повторы/однокоренные
                if len(low) >= MIN_WORD_LEN and re.match(r'^[а-яёа-яё]+$', low):
                    lem = get_lemma(low)
                    if lem in lemma_to_pid:
                        pid = lemma_to_pid[lem]
                        run.font.color.rgb = pair_rgbs[pid % len(pair_rgbs)]
                        run.bold = True

        # --- Отчёт в конце документа ---
        new_doc.add_page_break()

        header = new_doc.add_paragraph()
        header_run = header.add_run("═" * 50)
        header_run.font.size = Pt(10)

        title = new_doc.add_paragraph()
        title_run = title.add_run("ОТЧЁТ TAUTOLOGY HUNTER")
        title_run.bold = True
        title_run.font.size = Pt(14)

        source = new_doc.add_paragraph()
        source_run = source.add_run(f"Файл: {os.path.basename(input_path)}")
        source_run.font.size = Pt(10)

        from datetime import datetime
        date_para = new_doc.add_paragraph()
        date_run = date_para.add_run(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        date_run.font.size = Pt(10)

        # Считаем статистику
        repeats = [i for i in issues if i["type"] == "ПОВТОР"]
        cognates = [i for i in issues if i["type"] == "ОДНОКОР"]
        weeds_found = [i for i in issues if i["type"] == "СОРНЯК"]

        stats = new_doc.add_paragraph()
        stats_run = stats.add_run(
            f"\nПовторов: {len(repeats)} | "
            f"Однокоренных: {len(cognates)} | "
            f"Сорняков: {len(weeds_found)}"
        )
        stats_run.bold = True
        stats_run.font.size = Pt(11)

        sep = new_doc.add_paragraph()
        sep.add_run("─" * 50).font.size = Pt(10)

        # Повторы
        if repeats:
            h = new_doc.add_paragraph()
            hr = h.add_run("ПОВТОРЫ")
            hr.bold = True
            hr.font.size = Pt(12)
            hr.font.color.rgb = hex_to_rgb("#CC0000")

            for iss in repeats:
                p = new_doc.add_paragraph()
                r = p.add_run(
                    f"«{iss['w1']}» ↔ «{iss['w2']}» ({iss['info']}) "
                    f"— предл. {iss['s1']} и {iss['s2']}"
                )
                r.font.size = Pt(10)
                color_idx = iss["pair_id"] % len(PAIR_COLORS)
                r.font.color.rgb = pair_rgbs[color_idx]

                ctx = p.add_run(f"\n   → {iss['ctx']}…")
                ctx.font.size = Pt(9)
                ctx.italic = True

        # Однокоренные
        if cognates:
            sep2 = new_doc.add_paragraph()
            sep2.add_run("─" * 50).font.size = Pt(10)

            h = new_doc.add_paragraph()
            hr = h.add_run("ОДНОКОРЕННЫЕ")
            hr.bold = True
            hr.font.size = Pt(12)
            hr.font.color.rgb = hex_to_rgb("#CC6600")

            for iss in cognates:
                p = new_doc.add_paragraph()
                r = p.add_run(
                    f"«{iss['w1']}» ↔ «{iss['w2']}» ({iss['info']}) "
                    f"— предл. {iss['s1']} и {iss['s2']}"
                )
                r.font.size = Pt(10)
                color_idx = iss["pair_id"] % len(PAIR_COLORS)
                r.font.color.rgb = pair_rgbs[color_idx]

                ctx = p.add_run(f"\n   → {iss['ctx']}…")
                ctx.font.size = Pt(9)
                ctx.italic = True

        # Сорняки
        if weeds_found:
            sep3 = new_doc.add_paragraph()
            sep3.add_run("─" * 50).font.size = Pt(10)

            h = new_doc.add_paragraph()
            hr = h.add_run("СОРНЯКИ")
            hr.bold = True
            hr.font.size = Pt(12)
            hr.font.color.rgb = weed_rgb

            for iss in weeds_found:
                p = new_doc.add_paragraph()
                r = p.add_run(f"«{iss['w1']}» — предл. {iss['s1']}")
                r.font.size = Pt(10)
                r.font.color.rgb = weed_rgb

                ctx = p.add_run(f"\n   → {iss['ctx']}…")
                ctx.font.size = Pt(9)
                ctx.italic = True

        # Сохраняем
        base, ext = os.path.splitext(input_path)
        output_path = f"{base}_проверка{ext}"

        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word", "*.docx")],
            title="Сохранить проверенный файл",
            initialfile=os.path.basename(output_path),
        )
        if not save_path:
            return

        try:
            new_doc.save(save_path)
            self.status_var.set(
                f"Готово! Повторов: {len(repeats)}, "
                f"однокор.: {len(cognates)}, "
                f"сорняков: {len(weeds_found)}"
            )
            if sys.platform == "win32":
                os.startfile(save_path)
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить:\n{e}")

if __name__ == "__main__":
    app = TautologyHunterApp()
    app.run()